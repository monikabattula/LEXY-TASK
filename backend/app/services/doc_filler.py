from pathlib import Path
from typing import Dict
from copy import deepcopy

from docx import Document
from docx.shared import RGBColor

from app.models.domain import Document as DocumentModel, Placeholder, Answer
from app.core.storage import document_upload_path, ensure_parent


def fill_document(
    doc_model: DocumentModel,
    answers: Dict[str, str],  # placeholder_id -> value
    placeholders: list[Placeholder],
    output_path: Path,
) -> Path:
    """
    Fill a .docx document with answers, replacing placeholder text.
    Uses fuzzy matching based on source_excerpt and paragraph locations.
    """
    # Load original document
    if not Path(doc_model.original_path).exists():
        raise FileNotFoundError(f"Original document not found: {doc_model.original_path}")
    
    doc = Document(doc_model.original_path)
    
    # Build mapping of placeholder_id -> value
    # Only use actual extracted values, not field names or descriptions
    placeholder_values = {}
    for p in placeholders:
        if p.id in answers:
            value = answers[p.id]
            if value and value.strip():
                cleaned_value = value.strip()
                
                # Reject values that are field names or descriptions
                # Check if it matches placeholder name, description, or common patterns
                value_lower = cleaned_value.lower()
                name_lower = p.name.lower().replace("_", " ")
                desc_lower = (p.description or "").lower()
                
                # Don't use if it's a field name pattern
                is_field_name = (
                    value_lower == name_lower or
                    value_lower == f"the {name_lower}" or
                    value_lower == desc_lower or
                    value_lower in ["the company name", "the investor name", "the date", "the amount"] or
                    value_lower.startswith("the ") and name_lower in value_lower
                )
                
                if not is_field_name:
                    placeholder_values[p.id] = cleaned_value
    
    # Replace placeholders in paragraphs
    # Strategy: match by paragraph_index and char_start/char_end if available
    # Otherwise, search for source_excerpt text
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text
        
        # Try to find placeholders in this paragraph
        for placeholder in placeholders:
            if placeholder.id not in placeholder_values:
                continue
            
            value = placeholder_values[placeholder.id]
            
            # If we have exact location, use it
            if (placeholder.paragraph_index == para_idx and 
                placeholder.char_start is not None and 
                placeholder.char_end is not None):
                # This is tricky - docx paragraph.text is read-only for replacement
                # We need to replace runs instead
                if placeholder.source_excerpt in para_text:
                    # Replace the excerpt text
                    para_text = para_text.replace(placeholder.source_excerpt, value, 1)
                    # Clear existing runs
                    paragraph.clear()
                    # Add new text
                    paragraph.add_run(para_text)
                continue
            
            # Otherwise, try fuzzy match on source_excerpt
            if placeholder.source_excerpt and placeholder.source_excerpt in para_text:
                para_text = para_text.replace(placeholder.source_excerpt, value, 1)
                paragraph.clear()
                paragraph.add_run(para_text)
                continue
            
            # Try matching by placeholder name if excerpt not found
            if placeholder.name and f"[{placeholder.name}]" in para_text:
                para_text = para_text.replace(f"[{placeholder.name}]", value, 1)
                paragraph.clear()
                paragraph.add_run(para_text)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for placeholder in placeholders:
                    if placeholder.id not in placeholder_values:
                        continue
                    value = placeholder_values[placeholder.id]
                    if placeholder.source_excerpt and placeholder.source_excerpt in cell_text:
                        # Replace in cell
                        cell.clear()
                        cell.add_paragraph(value)
    
    # Save filled document
    ensure_parent(output_path)
    doc.save(output_path)
    return output_path


def generate_html_preview(
    doc_model: DocumentModel,
    answers: Dict[str, str],
    placeholders: list[Placeholder],
    output_path: Path,
) -> Path:
    """
    Generate an HTML preview of the filled document with actual content.
    """
    from app.services.doc_parser import extract_text_from_docx
    
    # Extract document text
    doc_path = Path(doc_model.original_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"Original document not found: {doc_model.original_path}")
    
    full_text, paragraphs = extract_text_from_docx(doc_path)
    
    # Create mapping: placeholder_id -> value and placeholder_name -> value for fallback
    placeholder_map = {}
    for p in placeholders:
        if p.id in answers:
            placeholder_map[p.id] = answers[p.id]
            placeholder_map[p.name] = answers[p.id]
    
    # Build document content with replaced placeholders
    content_paragraphs = []
    for para_idx, (orig_idx, para_text) in enumerate(paragraphs):
        # Replace placeholders in this paragraph
        processed_text = para_text
        
        # Replace each placeholder
        for placeholder in placeholders:
            if placeholder.id not in placeholder_map:
                continue
            
            value = placeholder_map[placeholder.id]
            
            # Try to replace source_excerpt
            if placeholder.source_excerpt and placeholder.source_excerpt in processed_text:
                processed_text = processed_text.replace(
                    placeholder.source_excerpt,
                    f'<span class="filled-value">{value}</span>',
                    1
                )
            
            # Also try bracket notation
            bracket_pattern = f"[{placeholder.name.upper()}]"
            if bracket_pattern in processed_text:
                processed_text = processed_text.replace(
                    bracket_pattern,
                    f'<span class="filled-value">{value}</span>',
                    1
                )
            
            # Try with underscores
            if placeholder.source_excerpt and "_____" in placeholder.source_excerpt:
                # Replace underscore patterns
                import re
                underscore_pattern = r'_{3,}'
                if re.search(underscore_pattern, processed_text):
                    processed_text = re.sub(
                        underscore_pattern,
                        f'<span class="filled-value">{value}</span>',
                        processed_text,
                        count=1
                    )
        
        # Only add non-empty paragraphs
        if processed_text.strip():
            content_paragraphs.append(processed_text)
    
    # Build HTML
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{doc_model.filename} - Preview</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.8;
            color: #333;
            background-color: #fafafa;
        }}
        .document-container {{
            background-color: white;
            padding: 60px 80px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            min-height: 800px;
        }}
        h1 {{
            color: #1a1a1a;
            margin-bottom: 30px;
            font-size: 24px;
            text-align: center;
            font-weight: bold;
        }}
        p {{
            margin-bottom: 12px;
            text-align: justify;
        }}
        .filled-value {{
            background-color: #e3f2fd;
            padding: 2px 6px;
            border-radius: 3px;
            font-weight: 500;
            color: #1565c0;
            border-bottom: 2px solid #64b5f6;
        }}
        .header-info {{
            background-color: #f5f5f5;
            padding: 20px;
            margin-bottom: 30px;
            border-radius: 8px;
            border-left: 4px solid #2196f3;
        }}
        .header-info p {{
            margin: 5px 0;
            text-align: left;
        }}
        .header-info strong {{
            color: #424242;
        }}
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            .document-container {{
                box-shadow: none;
                padding: 40px;
            }}
        }}
    </style>
</head>
<body>
    <div class="header-info">
        <p><strong>Document:</strong> {doc_model.filename}</p>
        <p><strong>Status:</strong> Filled with {len(answers)} placeholder(s)</p>
    </div>
    <div class="document-container">
"""
    
    # Add all paragraphs
    for para_text in content_paragraphs:
        # Escape HTML but preserve our filled-value spans
        import html as html_module
        # Split by our spans, escape the rest
        parts = para_text.split('<span class="filled-value">')
        escaped_parts = []
        for i, part in enumerate(parts):
            if '</span>' in part:
                value_part, rest = part.split('</span>', 1)
                escaped_parts.append(f'<span class="filled-value">{html_module.escape(value_part)}</span>{html_module.escape(rest)}')
            else:
                escaped_parts.append(html_module.escape(part))
        
        final_text = ''.join(escaped_parts)
        
        # Convert double newlines to paragraph breaks
        if final_text.strip():
            html_content += f"        <p>{final_text}</p>\n"
    
    html_content += """    </div>
</body>
</html>
"""
    
    ensure_parent(output_path)
    output_path.write_text(html_content, encoding="utf-8")
    return output_path

